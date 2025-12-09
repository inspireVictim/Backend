#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для конвертации Markdown в DOCX и создания ZIP архива
"""

import re
import zipfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def parse_markdown_to_docx(md_file_path, docx_file_path):
    """Конвертирует Markdown файл в DOCX"""
    
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Читаем MD файл
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    code_language = ''
    
    while i < len(lines):
        line = lines[i].rstrip()
        original_line = line
        
        # Обработка блоков кода
        if line.strip().startswith('```'):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                code_language = line.strip()[3:].strip()
                i += 1
                continue
        
        if in_code_block:
            # Добавляем код как моноширинный текст
            p = doc.add_paragraph(line, style='No Spacing')
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            i += 1
            continue
        
        # Пропускаем пустые строки
        if not line.strip():
            i += 1
            continue
        
        # Заголовки
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:].strip(), level=4)
        # Горизонтальная линия
        elif line.strip() == '---':
            p = doc.add_paragraph('─' * 50)
        # Список (маркированный)
        elif line.strip().startswith('- '):
            text = line[2:].strip()
            # Обрабатываем форматирование в списке
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
        # Список (нумерованный)
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
        # Обычный текст
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)
        
        i += 1
    
    # Сохраняем DOCX
    doc.save(docx_file_path)
    print(f"✓ DOCX файл создан: {docx_file_path}")

def add_formatted_text(paragraph, text):
    """Добавляет текст с форматированием (жирный, код, ссылки)"""
    # Обрабатываем комбинации форматирования
    # Сначала обрабатываем код (чтобы не конфликтовал с жирным)
    parts = re.split(r'(`[^`]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            # Inline код
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            # Обрабатываем жирный текст
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
            for bold_part in bold_parts:
                if bold_part.startswith('**') and bold_part.endswith('**'):
                    run = paragraph.add_run(bold_part[2:-2])
                    run.bold = True
                else:
                    # Обрабатываем ссылки [текст](url)
                    link_parts = re.split(r'(\[[^\]]+\]\([^\)]+\))', bold_part)
                    for link_part in link_parts:
                        match = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', link_part)
                        if match:
                            run = paragraph.add_run(match.group(1))
                            # В DOCX ссылки требуют специальной обработки, упростим
                        else:
                            paragraph.add_run(link_part)

def create_zip_archive(docx_file_path, zip_file_path):
    """Создает ZIP архив с DOCX файлом"""
    
    with zipfile.ZipFile(zip_file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        zipf.write(docx_file_path, Path(docx_file_path).name)
    
    print(f"✓ ZIP архив создан: {zip_file_path}")

def main():
    """Главная функция"""
    
    # Пути к файлам
    md_file = Path('BACKEND_DEVELOPER_KNOWLEDGE.md')
    docx_file = Path('BACKEND_DEVELOPER_KNOWLEDGE.docx')
    zip_file = Path('BACKEND_DEVELOPER_KNOWLEDGE.zip')
    
    if not md_file.exists():
        print(f"❌ Ошибка: файл {md_file} не найден!")
        return
    
    try:
        # Конвертируем MD в DOCX
        parse_markdown_to_docx(md_file, docx_file)
        
        # Создаем ZIP архив
        create_zip_archive(docx_file, zip_file)
        
        print("\n✅ Конвертация завершена успешно!")
        print(f"📄 DOCX файл: {docx_file.absolute()}")
        print(f"📦 ZIP архив: {zip_file.absolute()}")
        
    except ImportError as e:
        print("❌ Ошибка: не установлена библиотека python-docx")
        print("Установите её командой: pip install python-docx")
    except Exception as e:
        print(f"❌ Ошибка: {e}")

if __name__ == '__main__':
    main()

